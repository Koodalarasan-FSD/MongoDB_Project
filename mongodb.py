from flask import Flask,jsonify,render_template,request,url_for,session,redirect,flash,send_file
from pymongo import MongoClient
from pymongo.server_api import ServerApi


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

import os


app=Flask(__name__)

# You should use a strong, random secret key
app.secret_key='528a822469db75086fe145b84215582f'

"""
def set_run_format(run):
    run.bold=True
    run.font.color.rgb=RGBColor(0,0,0)  # RGB color for black
"""

def set_cell_format(cell,is_header=False):
    cell.paragraphs[0].runs[0].font.bold=is_header
    cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0,0,0)  if is_header else RGBColor(0, 0, 0)
    cell.paragraphs[0].runs[0].font.size=Pt(10)
    #cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    cell.paragraphs[0].paragraph_format.space_before = Pt(top)
    cell.paragraphs[0].paragraph_format.space_after = Pt(bottom)
    cell.paragraphs[0].paragraph_format.left_indent = Pt(left)
    cell.paragraphs[0].paragraph_format.right_indent = Pt(right)




@app.route('/')
def index():
    return render_template('mongodb.html')

@app.route('/get_properties',methods=['GET','POST'])
def get_properties():

    
    MONGODB_LINK=request.form['mongodblink']
    MONGODB_DATABASENAME=request.form['mongodb_databasename']
    MONGODB_COLLECTION_NAME=request.form['mongodb_Collection_name']

    print(MONGODB_LINK)
    print(MONGODB_DATABASENAME)
    print(MONGODB_COLLECTION_NAME)

    # Connect to MongoDB
    client=MongoClient(MONGODB_LINK)
    db=client[MONGODB_DATABASENAME]
    collection=db[MONGODB_COLLECTION_NAME]

    # Retrieve data from MongoDB
    data=list(collection.find())    # Convert the cursor to a list for easy iteration

 
    # Create a new Word document
    doc=Document()

    # Add a table to the document
    table=doc.add_table(rows=1,cols=len(data[0].keys()),style='Table Grid')

    # Add the header row to the table
    header_row=table.rows[0]
    for col_num, key in enumerate(data[0].keys()):
        cell=header_row.cells[col_num]
        cell.text=key
        #set_run_format(cell.paragraphs[0].runs[0])  # Apply formatting to the first run in the cell
        set_cell_format(cell, is_header=True)  # Apply formatting to the header cell
        set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the header row



    # Iterate through MongoDB data and write to the document
    for record in data:
        row_cells=table.add_row().cells
        for col_num,value in enumerate(record.values()):
            cell=row_cells[col_num]
            cell.text=str(value)
            set_cell_format(cell)  # Apply formatting to the header cell
            set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the data rows

            

    # Specify the full path for saving the document
    output_path = os.path.join(os.getcwd(), 'output_document.docx')

        
    # Save the document
    doc.save(output_path)
            
    # Set a session variable to indicate that the document is ready for download
    #session['document_ready']=True

    # Redirect to the index page
    #return redirect(url_for('index'))

    # Inform the user about the successful operation
    return f"Success: Document created successfully at {output_path}"
    
"""
@app.route('/download_file')
def download_file():

    # Remove the session variable after download
    session.pop('document_ready',None)

    # Specify the full path of the document
    file_path=os.path.join(os.getcwd(),'output_document.docx')

    return send_file(file_path, as_attachment=True)
"""


if __name__=="__main__":
    app.run()