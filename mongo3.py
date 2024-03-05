from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt,RGBColor
import os
import json

app = Flask(__name__)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    cell.paragraphs[0].paragraph_format.space_before = Pt(top)
    cell.paragraphs[0].paragraph_format.space_after = Pt(bottom)
    cell.paragraphs[0].paragraph_format.left_indent = Pt(left)
    cell.paragraphs[0].paragraph_format.right_indent = Pt(right)


@app.route('/')
def index():
    return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta http-equiv="X-UA-Compatible" content="IE=edge">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>JSON to Word</title>
        </head>
        <body>
            <h1>JSON to Word Converter</h1>
            <form action="/generate_word" method="post" enctype="multipart/form-data">
                <label for="jsonFile">Upload JSON Text File (.txt):</label>
                <input type="file" id="jsonFile" name="jsonFile" accept=".txt" required><br>
                <button type="submit">Generate Word</button>
            </form>
        </body>
        </html>



    """

@app.route('/generate_word', methods=['POST'])
def generate_word():
    # Check if the file is present in the request
    if 'jsonFile' not in request.files:
        return "No file provided"

    file = request.files['jsonFile']

    # Check if the file has a .txt extension
    if not file.filename.endswith('.txt'):
        return "Invalid file format. Please upload a text file with a .txt extension."

    try:
        # Read the JSON data from the file
        json_data = json.load(file)
    except json.JSONDecodeError as e:
        return f"Error parsing JSON: {str(e)}"

    # Create a new Word document
    doc = Document()

    # Get the keys from the JSON data
    keys=list(json_data.keys())

    # Add a table to the Word document with dynamic headings
    table=doc.add_table(rows=1,cols=len(keys))
    table.style="TableGrid"
    

    # Populate the table with key-value pairs
    for col,key in enumerate(keys):
        cell = table.cell(0, col)
        cell.text = key
        cell.paragraphs[0].runs[0].bold = True  # Make the heading bold
        cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0,0,0) # Makes header in black colour
        cell.paragraphs[0].runs[0].font.size=Pt(10) 
        set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the header row

    # Populate the table with key-value pairs
    row_cells=table.add_row().cells
    for col,key in enumerate(keys):
        cell=row_cells[col]
        cell.text=str(json_data[key])
        set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the row

  
    # Save the Word document in the current working directory
    save_path = os.path.join(os.getcwd(), 'output.docx')
    doc.save(save_path)

    # Send the Word document as a response
    return send_file(save_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
