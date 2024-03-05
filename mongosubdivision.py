from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
import os
import json

app = Flask(__name__)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    cell.paragraphs[0].paragraph_format.space_before = Pt(top)
    cell.paragraphs[0].paragraph_format.space_after = Pt(bottom)
    cell.paragraphs[0].paragraph_format.left_indent = Pt(left)
    cell.paragraphs[0].paragraph_format.right_indent = Pt(right)

def process_json(data, table, row_index=0):
    for key, value in data.items():
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    row_index = process_json(item, table, row_index)
                else:
                    # Add a new row for each item in the list
                    table.add_row()
                    row_index += 1
                    for col, col_key in enumerate(table.columns):
                        if col_key == key:
                            cell = table.cell(row_index, col)
                            cell.text = str(item)
        elif isinstance(value, dict):
            row_index = process_json(value, table, row_index)
        else:
            # Add a new row for non-list, non-dictionary values
            table.add_row()
            row_index += 1
            for col, col_key in enumerate(table.columns):
                if col_key == key:
                    cell = table.cell(row_index, col)
                    cell.text = str(value)

    return row_index

@app.route('/')
def index():
    return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta http-equiv="X-UA-Compatible" content="IE=edge">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>JSON to Word</title>
        </head>
        <body>
            <h1>JSON to Word Converter</h1>
            <form action="/generate_word" method="post" enctype="multipart/form-data">
                <label for="jsonFile">Upload JSON Text File (.txt):</label>
                <input type="file" id="jsonFile" name="jsonFile" accept=".txt" required><br>
                <button type="submit">Generate Word</button>
            </form>
        </body>
        </html>
    """

@app.route('/generate_word', methods=['POST'])
def generate_word():
    if 'jsonFile' not in request.files:
        return "No file provided"

    file = request.files['jsonFile']

    if not file.filename.endswith('.txt'):
        return "Invalid file format. Please upload a text file with a .txt extension."

    try:
        json_data = json.load(file)
    except json.JSONDecodeError as e:
        return f"Error parsing JSON: {str(e)}"

    doc = Document()
    keys = list(json_data.keys())
    table = doc.add_table(rows=1, cols=len(keys))
    table.style = "TableGrid"

    # Set headings for the first row
    for col, key in enumerate(keys):
        cell = table.cell(0, col)
        cell.text = key
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_margins(cell, top=5, bottom=5, left=5, right=5)

    # Populate the table with key-value pairs and handle nested structures
    process_json(json_data, table)

    save_path = os.path.join(os.getcwd(), 'output.docx')
    doc.save(save_path)

    return send_file(save_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
